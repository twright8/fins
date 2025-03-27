"""
Document loader module for handling various document types.
Adapted from the reference implementation with modifications for
synchronous subprocess execution.
"""
import os
import sys
import hashlib
import time
from pathlib import Path
import fitz  # PyMuPDF
import docx
import pandas as pd
import pytesseract
from PIL import Image
import io
import json
import concurrent.futures
import multiprocessing
from multiprocessing import Lock

# Add parent directory to sys.path to enable imports from project root
sys.path.append(str(Path(__file__).resolve().parent.parent.parent))
from src.core.config import CONFIG
from src.utils.logger import setup_logger
from src.utils.resource_monitor import log_memory_usage

logger = setup_logger(__name__)

# Create a lock for thread-safe access to the OCR cache
ocr_cache_lock = Lock()

class DocumentLoader:
    """Document loader that handles various file formats."""
    
    def __init__(self, temp_dir=None, ocr_cache_dir=None, status_queue=None):
        """
        Initialize document loader.
        
        Args:
            temp_dir (str, optional): Directory for temporary files. 
            ocr_cache_dir (str, optional): Directory for OCR cache.
            status_queue (Queue, optional): Queue for status updates.
        """
        base_dir = Path(__file__).resolve().parent.parent.parent
        self.temp_dir = Path(temp_dir) if temp_dir else base_dir / "temp"
        self.ocr_cache_dir = Path(ocr_cache_dir) if ocr_cache_dir else base_dir / "data" / "ocr_cache"
        self.status_queue = status_queue
        
        # Number of parallel OCR jobs
        self.ocr_parallel_jobs = CONFIG["document_processing"]["ocr_parallel_jobs"]
        
        # Ensure directories exist
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.ocr_cache_dir, exist_ok=True)
        
        logger.info(f"Document loader initialized with temp_dir={self.temp_dir}, "
                   f"ocr_cache_dir={self.ocr_cache_dir}, "
                   f"ocr_parallel_jobs={self.ocr_parallel_jobs}")
        
        if self.status_queue:
            self.status_queue.put(('status', 'Document loader initialized'))
    
    def _update_status(self, status, progress=None):
        """
        Update status via queue if available.
        
        Args:
            status (str): Status message
            progress (float, optional): Progress value between 0 and 1
        """
        if self.status_queue:
            if progress is not None:
                self.status_queue.put(('progress', progress, status))
            else:
                self.status_queue.put(('status', status))
        
        # Always log status
        logger.info(status)
    
    def _get_cache_key(self, image_data):
        """
        Generate a cache key from image data.
        
        Args:
            image_data (bytes): Image data
            
        Returns:
            str: Cache key
        """
        return hashlib.md5(image_data).hexdigest()
    
    def _get_cached_ocr(self, cache_key):
        """
        Get OCR result from cache.
        
        Args:
            cache_key (str): Cache key
            
        Returns:
            str or None: OCR text if cached, None otherwise
        """
        cache_file = self.ocr_cache_dir / f"{cache_key}.json"
        if cache_file.exists():
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cached_data = json.load(f)
                    logger.info(f"OCR cache hit: {cache_key}")
                    return cached_data.get('text', '')
            except Exception as e:
                logger.warning(f"Error reading OCR cache: {e}")
        
        return None
    
    def _save_to_ocr_cache(self, cache_key, text):
        """
        Save OCR result to cache.
        
        Args:
            cache_key (str): Cache key
            text (str): OCR text
        """
        cache_file = self.ocr_cache_dir / f"{cache_key}.json"
        try:
            with ocr_cache_lock:
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump({'text': text}, f)
            logger.info(f"Saved to OCR cache: {cache_key}")
        except Exception as e:
            logger.warning(f"Error saving to OCR cache: {e}")
    
    def load_document(self, file_path):
        """
        Load a document from the given file path.
        
        Args:
            file_path (str): Path to the document
            
        Returns:
            dict: Document data including metadata and text content
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self._update_status(f"Loading document: {file_path}")
        
        # Extract file extension and call appropriate handler
        file_extension = file_path.suffix.lower()
        
        handlers = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # May need additional handling for older .doc files
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv,
            '.txt': self._process_text,
        }
        
        if file_extension in handlers:
            result = handlers[file_extension](file_path)
            log_memory_usage(logger)
            return result
        else:
            error_msg = f"Unsupported file format: {file_extension}"
            self._update_status(error_msg)
            raise ValueError(error_msg)
    
    def _process_pdf(self, file_path):
        """
        Process a PDF file.
        
        Args:
            file_path (Path): Path to the PDF file
            
        Returns:
            dict: Document data including metadata and text content
        """
        self._update_status(f"Processing PDF file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'pdf',
            'content': [],
            'metadata': {},
            'images': []
        }
        
        try:
            # Open the PDF
            pdf_document = fitz.open(file_path)
            
            # Extract metadata
            doc_data['metadata'] = {
                'title': pdf_document.metadata.get('title', ''),
                'author': pdf_document.metadata.get('author', ''),
                'creation_date': pdf_document.metadata.get('creationDate', ''),
                'modification_date': pdf_document.metadata.get('modDate', ''),
                'page_count': len(pdf_document)
            }
            
            # Create a unified approach to ensure we get all text from a PDF
            # First, extract the text normally from all pages
            page_contents = []
            total_pages = len(pdf_document)
            
            for page_num, page in enumerate(pdf_document):
                progress = 0.1 + ((page_num / total_pages) * 0.4)  # Progress from 10% to 50%
                self._update_status(f"Extracting text from page {page_num+1}/{total_pages}", progress)
                
                page_text = page.get_text()
                page_contents.append({
                    'page_num': page_num + 1,
                    'text': page_text,
                    'page': page,
                    'needs_ocr': len(page_text.strip()) < 100  # Flag pages with little text
                })
            
            # Collect all pages that need OCR and all images that need OCR
            ocr_tasks = []
            
            # Pages that need OCR
            for page_data in page_contents:
                if page_data['needs_ocr']:
                    ocr_tasks.append(('page', page_data['page'], page_data['page_num']))
            
            # Images from the PDF
            for page_num, page in enumerate(pdf_document):
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Only process images of reasonable size (to avoid tiny icons)
                    img = Image.open(io.BytesIO(image_bytes))
                    width, height = img.size
                    if width > 100 and height > 100:  # Skip very small images
                        ocr_tasks.append(('image', image_bytes, (page_num + 1, img_index)))
            
            self._update_status(f"Collected {len(ocr_tasks)} OCR tasks", 0.5)
            
            # Process OCR tasks in parallel
            ocr_results = {}
            if ocr_tasks:
                total_ocr_tasks = len(ocr_tasks)
                completed_tasks = 0
                
                with concurrent.futures.ThreadPoolExecutor(max_workers=self.ocr_parallel_jobs) as executor:
                    future_to_task = {
                        executor.submit(self._perform_ocr_task, task_type, content, identifier): 
                        (task_type, identifier) 
                        for task_type, content, identifier in ocr_tasks
                    }
                    
                    for future in concurrent.futures.as_completed(future_to_task):
                        task_type, identifier, ocr_text = future.result()
                        if ocr_text:
                            ocr_results[(task_type, identifier)] = ocr_text
                        
                        completed_tasks += 1
                        progress = 0.5 + ((completed_tasks / total_ocr_tasks) * 0.4)  # Progress from 50% to 90%
                        self._update_status(f"OCR progress: {completed_tasks}/{total_ocr_tasks}", progress)
            
            # Update page contents with OCR results
            self._update_status("Finalizing document processing", 0.9)
            
            for page_data in page_contents:
                page_num = page_data['page_num']
                if page_data['needs_ocr'] and ('page', page_data['page_num']) in ocr_results:
                    page_data['text'] = ocr_results[('page', page_data['page_num'])]
                
                # Add image OCR text if available
                for (task_type, identifier), ocr_text in ocr_results.items():
                    if task_type == 'image' and identifier[0] == page_num:
                        img_index = identifier[1]
                        
                        # Add to images list
                        doc_data['images'].append({
                            'page': page_num - 1,  # 0-indexed
                            'index': img_index,
                            'ocr_text': ocr_text
                        })
                        
                        # Augment page text with image OCR
                        page_data['text'] += f"\n\n[Image OCR Text: {ocr_text}]"
            
            # Create final content
            for page_data in page_contents:
                doc_data['content'].append({
                    'page_num': page_data['page_num'],
                    'text': page_data['text']
                })
            
            # Close the document
            pdf_document.close()
            
            self._update_status(f"PDF processing complete: {file_path} - {len(doc_data['content'])} pages", 1.0)
            return doc_data
            
        except Exception as e:
            error_msg = f"Error processing PDF file {file_path}: {e}"
            self._update_status(error_msg)
            logger.error(error_msg)
            raise
    
    def _perform_ocr_task(self, task_type, content, identifier):
        """
        Perform OCR on a task.
        
        Args:
            task_type (str): 'page' or 'image'
            content: Page object or image bytes
            identifier: Page number or (page_num, img_index) tuple
            
        Returns:
            tuple: (task_type, identifier, ocr_text)
        """
        try:
            if task_type == 'page':
                page = content
                page_num = identifier
                
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))  # Lower resolution
                img_bytes = pix.tobytes()
                
                # Try cache first
                cache_key = self._get_cache_key(img_bytes)
                cached_text = self._get_cached_ocr(cache_key)
                
                if cached_text is not None:
                    return (task_type, page_num, cached_text)
                
                # Perform OCR
                image = Image.open(io.BytesIO(img_bytes))
                ocr_text = pytesseract.image_to_string(image)
                image.close()
                
                # Cache the result
                self._save_to_ocr_cache(cache_key, ocr_text)
                
                return (task_type, page_num, ocr_text)
                
            elif task_type == 'image':
                image_bytes = content
                page_img_index = identifier
                
                # Try cache first
                cache_key = self._get_cache_key(image_bytes)
                cached_text = self._get_cached_ocr(cache_key)
                
                if cached_text is not None:
                    return (task_type, page_img_index, cached_text)
                
                # Perform OCR
                image = Image.open(io.BytesIO(image_bytes))
                ocr_text = pytesseract.image_to_string(image)
                image.close()
                
                # Only store if it found text
                if len(ocr_text.strip()) > 10:
                    # Cache the result
                    self._save_to_ocr_cache(cache_key, ocr_text)
                    return (task_type, page_img_index, ocr_text)
                
                return (task_type, page_img_index, '')
            
            return (task_type, identifier, '')
            
        except Exception as e:
            logger.error(f"OCR error for {task_type} {identifier}: {e}")
            return (task_type, identifier, '')
    
    def _process_docx(self, file_path):
        """
        Process a DOCX file.
        
        Args:
            file_path (Path): Path to the DOCX file
            
        Returns:
            dict: Document data including metadata and text content
        """
        self._update_status(f"Processing DOCX file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'docx',
            'content': [],
            'metadata': {}
        }
        
        try:
            # Open the document
            document = docx.Document(file_path)
            
            # Extract metadata
            core_properties = document.core_properties
            doc_data['metadata'] = {
                'title': core_properties.title if hasattr(core_properties, 'title') else '',
                'author': core_properties.author if hasattr(core_properties, 'author') else '',
                'created': str(core_properties.created) if hasattr(core_properties, 'created') else '',
                'modified': str(core_properties.modified) if hasattr(core_properties, 'modified') else '',
                'paragraph_count': len(document.paragraphs),
                'table_count': len(document.tables)
            }
            
            self._update_status(f"Extracting content from DOCX", 0.3)
            
            # Extract content from paragraphs
            full_text = []
            for para in document.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            self._update_status(f"Extracting tables from DOCX", 0.6)
            
            # Extract content from tables
            table_texts = []
            for table in document.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                table_texts.append("\n".join(table_text))
            
            # Combine all text
            doc_data['content'].append({
                'text': "\n\n".join(full_text + ["\n\nTABLES:\n"] + table_texts if table_texts else full_text)
            })
            
            self._update_status(f"DOCX processing complete: {file_path}", 1.0)
            return doc_data
            
        except Exception as e:
            error_msg = f"Error processing DOCX file {file_path}: {e}"
            self._update_status(error_msg)
            logger.error(error_msg)
            raise
    
    def _process_excel(self, file_path):
        """
        Process an Excel file.
        
        Args:
            file_path (Path): Path to the Excel file
            
        Returns:
            dict: Document data including metadata and text content
        """
        self._update_status(f"Processing Excel file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'excel',
            'content': [],
            'metadata': {},
            'sheets': []
        }
        
        try:
            # Read Excel file with pandas
            excel_file = pd.ExcelFile(file_path)
            
            # Extract metadata
            doc_data['metadata'] = {
                'sheet_names': excel_file.sheet_names,
                'sheet_count': len(excel_file.sheet_names)
            }
            
            # Process each sheet
            all_text_content = []
            sheet_count = len(excel_file.sheet_names)
            
            for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                progress = 0.2 + ((sheet_idx / sheet_count) * 0.7)  # Progress from 20% to 90%
                self._update_status(f"Processing sheet {sheet_idx+1}/{sheet_count}: {sheet_name}", progress)
                
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Convert sheet to structured text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                sheet_text += f"Column headers: {', '.join(str(col) for col in df.columns)}\n\n"
                
                # Add sample data (first 10 rows)
                sample_rows = min(10, len(df))
                if sample_rows > 0:
                    sheet_text += "Sample data:\n"
                    sample_df = df.head(sample_rows)
                    sheet_text += sample_df.to_string() + "\n\n"
                
                # Save the full sheet data for potential further processing
                doc_data['sheets'].append({
                    'name': sheet_name,
                    'dataframe': df.to_dict(orient='records')
                })
                
                all_text_content.append(sheet_text)
            
            # Combine all sheets into one content item
            doc_data['content'].append({
                'text': "\n\n".join(all_text_content)
            })
            
            self._update_status(f"Excel processing complete: {file_path} - {len(doc_data['sheets'])} sheets", 1.0)
            return doc_data
            
        except Exception as e:
            error_msg = f"Error processing Excel file {file_path}: {e}"
            self._update_status(error_msg)
            logger.error(error_msg)
            raise
    
    def _process_csv(self, file_path):
        """
        Process a CSV file.
        
        Args:
            file_path (Path): Path to the CSV file
            
        Returns:
            dict: Document data including metadata and text content
        """
        self._update_status(f"Processing CSV file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'csv',
            'content': [],
            'metadata': {}
        }
        
        try:
            # Read CSV file
            self._update_status(f"Reading CSV file", 0.3)
            df = pd.read_csv(file_path)
            
            # Extract metadata
            doc_data['metadata'] = {
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': list(df.columns)
            }
            
            self._update_status(f"Converting CSV to structured text", 0.6)
            
            # Convert to structured text
            text_content = f"CSV File: {file_path.name}\n"
            text_content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
            text_content += f"Column headers: {', '.join(str(col) for col in df.columns)}\n\n"
            
            # Add sample data (first 10 rows)
            sample_rows = min(10, len(df))
            if sample_rows > 0:
                text_content += "Sample data:\n"
                sample_df = df.head(sample_rows)
                text_content += sample_df.to_string() + "\n\n"
            
            # Store the full dataframe for potential further processing
            doc_data['dataframe'] = df.to_dict(orient='records')
            
            # Add content for text extraction
            doc_data['content'].append({
                'text': text_content
            })
            
            self._update_status(f"CSV processing complete: {file_path} - {doc_data['metadata']['row_count']} rows", 1.0)
            return doc_data
            
        except Exception as e:
            error_msg = f"Error processing CSV file {file_path}: {e}"
            self._update_status(error_msg)
            logger.error(error_msg)
            raise
    
    def _process_text(self, file_path):
        """
        Process a plain text file.
        
        Args:
            file_path (Path): Path to the text file
            
        Returns:
            dict: Document data including metadata and text content
        """
        self._update_status(f"Processing text file: {file_path}")
        
        doc_data = {
            'document_id': str(file_path.stem),
            'file_name': file_path.name,
            'file_type': 'text',
            'content': [],
            'metadata': {}
        }
        
        try:
            # Read text file
            self._update_status(f"Reading text file", 0.5)
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text_content = f.read()
            
            # Extract basic metadata
            doc_data['metadata'] = {
                'file_size': os.path.getsize(file_path),
                'line_count': text_content.count('\n') + 1,
                'character_count': len(text_content)
            }
            
            # Add content
            doc_data['content'].append({
                'text': text_content
            })
            
            self._update_status(f"Text file processing complete: {file_path} - {doc_data['metadata']['character_count']} characters", 1.0)
            return doc_data
            
        except Exception as e:
            error_msg = f"Error processing text file {file_path}: {e}"
            self._update_status(error_msg)
            logger.error(error_msg)
            raise
